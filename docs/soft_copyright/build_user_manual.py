from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "soft_copyright" / "WS63激光打标控制软件用户手册_V1.0.docx"

FONT_CN = "Microsoft YaHei"
FONT_MONO = "Consolas"
COLOR_BLUE = "1F4E79"
COLOR_LIGHT_BLUE = "D9EAF7"
COLOR_GRAY = "F2F2F2"
COLOR_WARN = "FFF2CC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name=FONT_CN, size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_font(paragraph, name=FONT_CN, size=10.5):
    for run in paragraph.runs:
        set_run_font(run, name=name, size=size)


def add_para(doc, text="", style=None, align=None, spacing_after=4, first_line=False):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = 1.25
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def add_title(doc, title, subtitle):
    for _ in range(5):
        add_para(doc)
    p = add_para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=16)
    r = p.runs[0]
    set_run_font(r, size=22, bold=True, color=COLOR_BLUE)
    p = add_para(doc, subtitle, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=32)
    set_run_font(p.runs[0], size=14, bold=True)

    rows = [
        ("软件全称", "基于海思WS63的可手持式无线激光振镜打标机控制软件"),
        ("软件简称", "WS63激光打标控制软件"),
        ("版本号", "V1.0"),
        ("文档类型", "用户手册"),
        ("适用平台", "海思WS63 LiteOS嵌入式平台"),
    ]
    table = add_table(doc, ["项目", "内容"], rows, widths=[3.2, 11.8])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_para(doc)
    p = add_para(doc, "本文档用于说明软件的运行环境、功能特性、操作流程、安全注意事项及常见问题处理方法。", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(p.runs[0], size=10)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=15, bold=True, color=COLOR_BLUE)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(r, size=12, bold=True, color=COLOR_BLUE)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(5)
    else:
        set_run_font(r, size=10.5, bold=True)
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(3)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, size=10.5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.2


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, size=10.5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.2


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, name=FONT_MONO, size=9)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], COLOR_LIGHT_BLUE)
        set_cell_margins(hdr_cells[i])
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_para_font(p, size=10)
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            for p in cells[i].paragraphs:
                set_para_font(p, size=9.5)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
    if widths:
        for row in table.rows:
            for idx, width_cm in enumerate(widths):
                row.cells[idx].width = Cm(width_cm)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_note(doc, title, text, fill=COLOR_WARN):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
    p = cell.paragraphs[0]
    r = p.add_run(title + "：")
    set_run_font(r, size=10.5, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    p.paragraph_format.line_spacing = 1.2
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = FONT_CN
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(10.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("WS63激光打标控制软件用户手册 V1.0")
    set_run_font(r, size=9, color="666666")


def build_manual():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    add_title(doc, "基于海思WS63的可手持式无线激光振镜打标机控制软件", "用户手册 V1.0")

    add_heading(doc, "目录", 1)
    toc = [
        "1 文档说明",
        "2 软件概述",
        "3 运行环境",
        "4 功能说明",
        "5 硬件连接",
        "6 安装、编译与烧录",
        "7 LaserGRBL连接与参数配置",
        "8 基本操作流程",
        "9 矢量打标、灰度扫描与切割雕刻",
        "10 常用命令与诊断信息",
        "11 安全注意事项",
        "12 常见问题处理",
        "13 维护与校准",
        "附录 术语说明",
    ]
    add_bullets(doc, toc)
    doc.add_page_break()

    add_heading(doc, "1 文档说明", 1)
    add_para(
        doc,
        "本手册面向可手持式无线激光振镜打标机的软件使用、调试和维护人员，说明软件的基本功能、运行环境、连接方式、操作流程、常用命令、安全要求和故障处理方法。",
        first_line=True,
    )
    add_table(
        doc,
        ["项目", "说明"],
        [
            ("适用软件", "基于海思WS63的可手持式无线激光振镜打标机控制软件 V1.0"),
            ("适用硬件", "海思WS63开发板、DAC8562、XY振镜、PWM激光模块及相关供电/保护电路"),
            ("适用用户", "设备操作者、嵌入式调试人员、比赛展示与验收人员"),
            ("文档用途", "用于软件操作说明、交付说明及软件著作权登记文档材料"),
        ],
        widths=[3.5, 11.5],
    )

    add_heading(doc, "2 软件概述", 1)
    add_para(
        doc,
        "本软件运行于海思WS63嵌入式平台，接收上位机发送的G-code控制指令，完成激光振镜打标过程中的运动插补、队列调度、PWM激光功率控制、DAC坐标输出、状态反馈和异常保护等功能。",
        first_line=True,
    )
    add_bullets(
        doc,
        [
            "支持LaserGRBL等上位机通过串口发送G-code任务。",
            "支持G0/G1直线运动、M3/M4/M5激光控制、S功率值和F速度值解析。",
            "支持矢量轮廓打标、灰度扫描、矢量切割、浅雕和深雕等应用场景。",
            "支持运动队列缓冲与任务结束同步，降低丢图、提前返回和StopResponding风险。",
            "支持$D诊断输出，可查看队列、运动、激光、PWM和时序统计信息。",
        ],
    )

    add_heading(doc, "3 运行环境", 1)
    add_heading(doc, "3.1 软件开发环境", 2)
    add_table(
        doc,
        ["类别", "配置"],
        [
            ("主机系统", "Windows 11，配合WSL2 Ubuntu开发环境"),
            ("开发工具", "海思WS63 SDK、GCC工具链、CMake/Python构建脚本、VS Code或其他代码编辑器"),
            ("编程语言", "C语言"),
            ("上位机软件", "LaserGRBL或其他支持标准G-code输出的激光打标/雕刻软件"),
        ],
        widths=[4.0, 11.0],
    )
    add_heading(doc, "3.2 软件运行环境", 2)
    add_table(
        doc,
        ["类别", "配置"],
        [
            ("运行平台", "海思WS63 LiteOS嵌入式系统"),
            ("主控硬件", "海思WS63开发板"),
            ("执行器件", "DAC8562数模转换器、XY振镜、PWM激光模块"),
            ("通信接口", "UART串口；无线版本可扩展星闪/近距离无线通信链路"),
        ],
        widths=[4.0, 11.0],
    )

    add_heading(doc, "4 功能说明", 1)
    add_table(
        doc,
        ["功能模块", "功能说明"],
        [
            ("G-code解析", "解析上位机发送的G0、G1、M3、M4、M5、S、F、$类命令，转换为内部运动与激光控制指令。"),
            ("运动控制", "根据目标坐标和速度进行直线插补，输出X/Y轴DAC电压，驱动振镜完成轨迹扫描。"),
            ("激光控制", "根据M3/M4/M5和S值控制激光使能与PWM占空比，实现恒功率或动态功率打标。"),
            ("队列调度", "将上位机指令缓存到运动队列，保证连续任务稳定执行，并在任务结束时等待队列真正完成。"),
            ("状态反馈", "向上位机返回ok、版本信息、实时状态和$D诊断信息，便于观察任务运行状态。"),
            ("安全保护", "在复位、急停、任务结束和M5命令后执行关光处理，减少异常出光风险。"),
        ],
        widths=[3.6, 11.4],
    )

    add_heading(doc, "5 硬件连接", 1)
    add_note(doc, "安全提示", "激光设备存在灼伤、致盲和引燃风险。接线、调试、聚焦和打标前必须佩戴对应波长的防护眼镜，并确保激光路径内无人员和易燃物。")
    add_table(
        doc,
        ["功能", "WS63引脚", "连接对象", "说明"],
        [
            ("UART TX", "GPIO15 / UART1_TXD", "USB转串口RX", "上位机接收板端状态和ok响应"),
            ("UART RX", "GPIO16 / UART1_RXD", "USB转串口TX", "板端接收上位机G-code命令"),
            ("DAC SCK", "GPIO7 / SPI0_SCK", "DAC8562 SCLK", "DAC SPI时钟"),
            ("DAC MOSI", "GPIO9 / SPI0_OUT", "DAC8562 DIN", "DAC SPI数据输出"),
            ("DAC CS", "GPIO10 / GPIO输出", "DAC8562 CS", "DAC片选控制"),
            ("Laser PWM", "GPIO2 / PWM2", "激光驱动PWM输入", "控制激光功率或调制强度"),
        ],
        widths=[3.0, 4.0, 4.2, 3.8],
    )
    add_bullets(
        doc,
        [
            "所有模块应共地连接，避免DAC输出、PWM信号和串口信号参考电平不一致。",
            "激光模块供电应满足额定电流要求，建议使用独立稳压供电并增加保险、急停或硬件使能保护。",
            "振镜和DAC模拟输出线应尽量短，并远离高电流激光电源线，降低噪声耦合。",
        ],
    )

    add_heading(doc, "6 安装、编译与烧录", 1)
    add_heading(doc, "6.1 获取工程", 2)
    add_para(doc, "将工程源码放置到WSL2开发目录中，确认SDK路径、工具链路径和目标板型号与当前工程配置一致。", first_line=True)
    add_heading(doc, "6.2 编译固件", 2)
    add_para(doc, "在WSL2终端进入工程src目录后执行构建命令：")
    add_code(doc, "cd /root/fbb_ws63/src")
    add_code(doc, "python3 build.py ws63-liteos-app")
    add_para(doc, "编译完成后，固件包通常生成在以下目录：")
    add_code(doc, "/root/fbb_ws63/src/output/ws63/fwpkg/ws63-liteos-app/ws63-liteos-app_all.fwpkg")
    add_heading(doc, "6.3 烧录固件", 2)
    add_steps(
        doc,
        [
            "连接WS63开发板与PC，确认下载串口可被识别。",
            "打开海思烧录工具或工程配套烧录脚本，选择编译生成的fwpkg固件包。",
            "使开发板进入下载模式，开始烧录并等待完成。",
            "烧录完成后复位开发板，使用串口工具或LaserGRBL验证版本信息和响应状态。",
        ],
    )

    add_heading(doc, "7 LaserGRBL连接与参数配置", 1)
    add_steps(
        doc,
        [
            "通过USB转串口连接PC与WS63板端UART接口。",
            "打开LaserGRBL，选择对应COM口，波特率设置为115200。",
            "点击连接，观察控制台是否收到版本信息和ok响应。",
            "发送$I查看固件标识，发送$D查看当前运动、队列和PWM诊断信息。",
            "导入待打标图像或矢量文件，根据材料设置速度、功率、线密度和模式。",
        ],
    )
    add_table(
        doc,
        ["参数", "建议初始值", "说明"],
        [
            ("串口波特率", "115200", "与固件UART配置保持一致"),
            ("矢量打标速度", "F3000-F8000", "根据材料和线条效果调整"),
            ("灰度扫描速度", "F2000-F4000", "先用灰阶测试图确认功率响应"),
            ("灰度线密度", "6-10 line/mm", "线密度越高越细腻，但耗时更长"),
            ("激光功率", "S100-S800", "根据材料阈值测试，不建议直接满功率"),
        ],
        widths=[3.6, 3.8, 7.6],
    )

    add_heading(doc, "8 基本操作流程", 1)
    add_steps(
        doc,
        [
            "检查激光防护、供电、接线、固定支架和加工区域安全状态。",
            "连接上位机并确认串口通信正常。",
            "执行$H或手动回零，使软件坐标与实际打标区域一致。",
            "使用弱功率或边框预览功能确认图案位置和尺寸。",
            "根据材料选择矢量、灰度、切割、浅雕或深雕参数。",
            "开始任务后观察LaserGRBL进度、板端状态和实际打标效果。",
            "任务结束后确认激光已关闭，必要时发送M5或断开硬件使能。",
        ],
    )

    add_heading(doc, "9 矢量打标、灰度扫描与切割雕刻", 1)
    add_heading(doc, "9.1 矢量打标", 2)
    add_para(doc, "矢量打标适合线稿、轮廓、文字、Logo和二维码边线等内容。建议使用清晰的矢量文件或高对比度位图转矢量，避免过多短碎线段造成路径冗余。", first_line=True)
    add_bullets(
        doc,
        [
            "建议先使用较低功率进行边框预览，确认图案没有越界。",
            "线条发虚时可降低速度或提高功率；线条烧蚀过重时降低功率或提高速度。",
            "复杂矢量图应减少重复路径、重叠路径和极短线段。",
        ],
    )
    add_heading(doc, "9.2 灰度扫描", 2)
    add_para(doc, "灰度扫描用于照片、渐变图、灰阶图等连续明暗内容。该模式对材料阈值、焦距、线密度、速度和PWM功率响应较敏感，应先使用灰阶测试图确定可用功率范围。", first_line=True)
    add_table(
        doc,
        ["现象", "调整方向"],
        [
            ("整体偏浅", "降低速度或提高最大功率S值"),
            ("整体偏深", "提高速度或降低最大功率S值"),
            ("颗粒感重", "提高线密度、使用更平滑的图片、降低速度并确认焦距"),
            ("横线明显", "检查线密度、机械固定、PWM响应和材料表面一致性"),
            ("灰阶分不开", "用功率/速度矩阵测试材料阈值，缩小有效功率区间"),
        ],
        widths=[4.0, 11.0],
    )
    add_heading(doc, "9.3 矢量切割", 2)
    add_para(doc, "矢量切割适合薄纸、薄木片、胶片等材料的轮廓切割。切割通常需要低速、高功率或多遍重复加工。", first_line=True)
    add_bullets(
        doc,
        [
            "建议先在废料上测试材料切透阈值。",
            "优先使用多遍中等功率切割，减少一次高功率导致的烧焦和变形。",
            "切割任务必须加强通风、防火和固定，禁止无人值守。",
        ],
    )
    add_heading(doc, "9.4 浅雕与深雕", 2)
    add_para(doc, "浅雕用于表面浅层变色或轻微去除，深雕用于更明显的材料去除。当前软件可通过速度、功率、线密度和重复次数组合实现浅雕或深雕效果。", first_line=True)
    add_table(
        doc,
        ["目标", "推荐调节"],
        [
            ("浅雕", "中高速、中低功率、较少重复次数"),
            ("深雕", "低速、中高功率、增加重复次数，注意散热和烟尘"),
            ("边缘清晰", "降低图像噪声，使用矢量轮廓或提高线密度"),
            ("表面均匀", "稳定焦距，固定材料，控制每次加工能量一致"),
        ],
        widths=[4.0, 11.0],
    )

    add_heading(doc, "10 常用命令与诊断信息", 1)
    add_table(
        doc,
        ["命令", "功能说明"],
        [
            ("$I", "查询软件版本和固件标识"),
            ("$G", "查询当前运动模式和坐标状态"),
            ("$D", "输出诊断信息，包括队列、运动、激光、PWM和时序统计"),
            ("$H", "执行回零或坐标复位流程"),
            ("$FRAME", "执行边框预览或轮廓检查"),
            ("?", "查询实时状态"),
            ("G0 X.. Y..", "快速移动到指定坐标，通常不出光"),
            ("G1 X.. Y.. F..", "按指定速度执行直线插补运动"),
            ("M3 S..", "开启恒功率激光输出"),
            ("M4 S..", "开启动态功率激光输出"),
            ("M5", "关闭激光并等待任务完成后硬关光"),
        ],
        widths=[3.4, 11.6],
    )
    add_note(doc, "诊断说明", "$D中的queue、busy、laser、power、pclk、period、high、late、seg等字段用于判断运动队列、PWM配置、激光状态和实时执行质量。正常任务结束后应看到queue=0、busy=0、laser=0。", fill=COLOR_GRAY)

    add_heading(doc, "11 安全注意事项", 1)
    add_bullets(
        doc,
        [
            "任何调试、预览和打标操作都必须佩戴与激光波长匹配的防护眼镜。",
            "禁止将激光照射方向对准人员、动物、反光镜面、玻璃或未知反射物。",
            "打标区域应具备遮光、防火和通风条件，附近不得放置易燃易爆物品。",
            "首次加工新材料必须从低功率、低占空比、短时间测试开始。",
            "设备异常出光、冒烟、异味、振镜失控或通信中断时，应立即急停或断电。",
            "切割和深雕任务不允许无人值守。",
        ],
    )

    add_heading(doc, "12 常见问题处理", 1)
    add_table(
        doc,
        ["问题", "可能原因", "处理方法"],
        [
            ("无法连接上位机", "COM口错误、波特率错误、串口线接反", "确认COM口、115200波特率、TX/RX交叉连接和供电状态"),
            ("振镜不动作", "DAC接线异常、SPI未输出、坐标范围错误", "检查DAC供电、SCK/MOSI/CS接线，发送小范围G0/G1测试"),
            ("激光不出光", "PWM接线错误、S值过低、M5关闭状态", "检查GPIO2 PWM输入、M3 S值和激光驱动使能"),
            ("出现StopResponding", "任务结束等待过久、串口响应中断、上位机超时", "查看$D诊断信息，确认队列是否持续变化，必要时复位并降低任务复杂度"),
            ("图像末尾丢失", "上位机完成状态早于板端真实执行完成", "使用带任务结束同步的固件版本，M5后等待队列清空"),
            ("灰度颗粒感重", "图片不适合灰度、线密度不足、功率区间未校准", "使用灰阶/渐变测试图，调整F、S和line/mm参数"),
            ("切割不透", "能量不足、速度过快、焦距不准", "降低速度、提高功率或增加重复次数，并重新聚焦"),
        ],
        widths=[3.4, 4.6, 7.0],
    )

    add_heading(doc, "13 维护与校准", 1)
    add_bullets(
        doc,
        [
            "定期检查激光模块、振镜、DAC和开发板连接端子，避免松动导致异常轨迹或间歇出光。",
            "更换材料、改变焦距或移动设备位置后，应重新进行边框预览和灰阶测试。",
            "长期使用后应检查振镜零点、打标尺寸比例和X/Y方向是否存在偏移。",
            "保留稳定固件版本和测试参数记录，便于出现异常时快速回退和对比。",
            "无线版本开发完成后，应额外验证链路丢包、延迟、流控和断连关光保护。",
        ],
    )

    add_heading(doc, "附录 术语说明", 1)
    add_table(
        doc,
        ["术语", "说明"],
        [
            ("G-code", "数控设备常用的运动控制指令格式，用于描述坐标、速度、功率和动作。"),
            ("振镜", "通过高速偏转镜片改变激光光斑位置的扫描执行机构。"),
            ("PWM", "脉宽调制信号，用于控制激光模块输入功率或调制强度。"),
            ("DAC", "数模转换器，将数字坐标值转换为模拟电压驱动振镜。"),
            ("S值", "G-code中的激光功率参数，数值越大通常表示输出功率越高。"),
            ("F值", "G-code中的运动速度参数，单位通常为mm/min。"),
            ("line/mm", "灰度扫描线密度，表示每毫米扫描多少行。"),
        ],
        widths=[3.2, 11.8],
    )

    add_heading(doc, "版本记录", 1)
    add_table(
        doc,
        ["版本", "日期", "说明"],
        [
            ("V1.0", "2026年6月", "形成用户手册初版，覆盖有线单板打标软件的主要使用流程与安全说明。"),
        ],
        widths=[3.0, 3.5, 8.5],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_manual()
